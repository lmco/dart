# -*- coding: utf8 -*-

# Copyright 2017 Lockheed Martin Corporation
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#

import logging
import json
import copy
import os
import shutil
import re
import imghdr
from cStringIO import StringIO
import traceback
import io

from django.utils.timezone import localtime, now
from django.conf import settings
from django.utils.text import normalize_newlines

from docx import Document
from docx.shared import Inches, RGBColor
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError

from missions.models import Mission, DARTDynamicSettings, TestDetail
from .helpers.sorters import TestSortingHelper
from .helpers.formatters import standardize_report_output_field


logger = logging.getLogger(__name__)


class ReturnStatus(object):
    def __init__(self, success=True, message='', **kwargs):
        self.success = success
        self.message = message
        self.data = kwargs.get('data') or {}

    def to_json(self):
        return json.dumps(self.__dict__)

    def to_dict(self):
        return self.__dict__

    def __str__(self):
        return str(self.success) + ': ' + str(self.message)


def copy_table(document, table, cut=False):
    """
    Copy a table and append it to the end of the document.
    :param document: A python-docx document object
    :param table: A python-docx table object
    :param cut: Instead of duplicating a the table, move it to the end
                of the document
    :return: The newly created table instance
    """
    if cut:
        document._body._element._insert_tbl(table._tbl)
    else:
        document._body._element._insert_tbl(copy.deepcopy(table)._tbl)
    return document.tables[-1]


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def get_cleared_paragraph(cell):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    return paragraph


def generate_report_or_attachments(mission_id, zip_attachments=False):
    '''
    Generates the report docx or attachments zip.
    :param mission_id: The id of the mission
    :param zip_attachments: True to return a zip of attachments, False to get the docx report as an IOStream
    :return: Returns StringIO if returning a report, a zip object otherwise
    '''
    system_classification = DARTDynamicSettings.objects.get_as_object().system_classification
    system_classification_verbose = system_classification.verbose_legend
    system_classification_short = system_classification.short_legend

    mission = Mission.objects.get(id=mission_id)
    tests = TestSortingHelper.get_ordered_testdetails(
        mission_id=mission_id,
        reportable_tests_only=True)

    # Set some values we'll use throughout this section
    total_reportable_tests = len(tests)
    total_tests_with_findings = TestDetail.objects.filter(
        mission=mission,
        has_findings=True).count()
    total_tests_without_findings = TestDetail.objects.filter(
        mission=mission,
        has_findings=False).count()
    LIGHTEST_PERMISSIBLE_CLASSIFICATION_LABEL_COLOR = 0xbbbbbb
    DARKEN_OVERLY_LIGHT_CLASSIFICATION_LABEL_COLOR_BY = 0x444444
    mission_data_dir = None
    report_has_attachments = False

    def replace_document_slugs(doc):
        """Cycle through the runs in each paragraph in the template & replace handlebar slugs"""

        logger.debug('> replace_document_slugs')

        handlebar_slugs = {
            r'{{AREA}}': str(mission.business_area),
            r'{{MISSION}}': str(mission.mission_name),
            r'{{GENERATION_DATE}}': now().strftime('%x'),
            r'{{TOTAL_TESTS}}': str(total_reportable_tests),
            r'{{TESTS_WITH_FINDINGS}}': str(total_tests_with_findings),
            r'{{TESTS_WITHOUT_FINDINGS}}': str(total_tests_without_findings),
        }

        for p in doc.paragraphs:
            for r in p.runs:
                for pattern in handlebar_slugs.keys():
                    if re.search(pattern, r.text):
                        logger.debug('>> Replaced: {old} With: {new}'.format(
                            old=r.text.encode('utf-8'),
                            new=handlebar_slugs[pattern].encode('utf-8')
                            )
                        )
                        r.text = re.sub(pattern, handlebar_slugs[pattern], r.text)

    def get_or_create_mission_data_dir(mission_data_dir):
        if mission_data_dir is None :
            mission_data_dir = os.path.join(
                settings.BASE_DIR,
                'SUPPORTING_DATA_PACKAGE',
                str(mission.id) + "_" + str(now().strftime('%Y%m%d-%H%M%S')))
            if not os.path.isdir(mission_data_dir):
                os.makedirs(mission_data_dir)
        return mission_data_dir

    def add_to_data_dir(mission_data_dir, test_case_number, supporting_data):
        # If test case directory path is not created, create it
        path = os.path.join(mission_data_dir, str(test_case_number))

        # Shouldn't have race condition issues, but may need to add
        # try/catch if so later on
        if not os.path.isdir(path):
            os.makedirs(path)

        # Copy file to destination path
        shutil.copy(
            os.path.join(settings.MEDIA_ROOT, supporting_data.filename()),
            path
        )

    def prepend_classification(text):
        return '(' + system_classification_short + ') ' + text

    def portion_mark_and_insert(paragraphs, document):
        for paragraph in normalize_newlines(paragraphs).split('\n'):
            if len(paragraph) > 0:
                document.add_paragraph(prepend_classification(paragraph))

    # Load the template file to get the styles
    document = Document(settings.REPORT_TEMPLATE_PATH)

    # Get the table templates
    table_no_findings = document.tables[0]
    table_with_findings = document.tables[1]
    data_table = document.tables[2]

    # Set the classification legend color from the background color of the banner
    classification_style = document.styles['Table Classification']
    classification_font = classification_style.font

    # RGBColor doesn't handle shorthand hex codes, so let's just go ahead and expand it
    # if we come across a legacy or "misguided" entry
    if len(system_classification.background_color.hex_color_code) == 3:
        new_hex_code = ''
        for char in system_classification.background_color.hex_color_code:
            new_hex_code += char + char
        system_classification.background_color.hex_color_code = new_hex_code
        system_classification.background_color.save()

    if len(system_classification.text_color.hex_color_code) == 3:
        new_hex_code = ''
        for char in system_classification.text_color.hex_color_code:
            new_hex_code += char + char
        system_classification.text_color.hex_color_code = new_hex_code
        system_classification.text_color.save()

    classification_font.color.rgb = RGBColor.from_string(system_classification.get_report_label_color().hex_color_code)

    # Intro H1 and text
    document.add_heading('Introduction', level=1)
    portion_mark_and_insert(mission.introduction, document)

    # Scope H1 and text
    document.add_heading('Scope', level=1)
    portion_mark_and_insert(mission.scope, document)

    # Objectives H1 and text
    document.add_heading('Objectives', level=1)
    portion_mark_and_insert(mission.objectives, document)

    # Exec Summary H1 and text
    document.add_heading('Executive Summary', level=1)
    portion_mark_and_insert(mission.executive_summary, document)

    # Technical Assessment / Attack Architecture and text H1
    document.add_heading('Technical Assessment / Attack Architecture', level=1)
    portion_mark_and_insert(mission.technical_assessment_overview, document)

    # Technical Assessment / Test Cases and Results and loop
    document.add_heading('Technical Assessment / Test Cases and Results', level=1)

    # For each test, Test # - Objective  Attack Phase: H2

    """
    Note: If a test case is hidden from the report, the test_case number and the auto-generated paragraph
    number in Word won't align and it could cause confusion. The team may need to discuss how they want to
    handle this situation.
    """

    """
    WARNING

    Hidden test cases are NOT numbered in the report so the customer doesn't think something was
    accidentally left out. This will result in test case numbering differing between report and web when
    one or more test cases are hidden from the report.
    """

    test_case_number = 0
    tests_with_findings = 0
    tests_without_findings = 0

    # Let's just always craete a mission data dir when zipping attachments on report generation; this way
    # mission attachments will always return a zip file... it just may be empty - located here to account for times
    # when Get Attachments is actuated, but there are no test cases yet.
    if zip_attachments:
        mission_data_dir = get_or_create_mission_data_dir(mission_data_dir)

    for t in tests:

        if test_case_number > 0:
            document.add_page_break()

        test_case_number += 1

        if t.has_findings:
            test_title = "*"
            tests_with_findings += 1
        else:
            test_title = ""
            tests_without_findings += 1

        if t.enclave:
            test_title += "(%s) %s" % (
                t.enclave,
                t.test_objective
            )
        else:
            test_title += "%s" % (
                t.test_objective,
            )
        document.add_heading(test_title, level=2)

        # Duplicate one of the pre-made tables; if this is the last of a specific type of
        # test case (findings / no findings), use the cut operation to remove the blank
        # table.
        #TODO: convert to logging: print("T w/ F: {0}\nTotal: {1}\nT w/o F: {2}\nTotal: {3}".format(tests_with_findings, total_tests_with_findings, tests_without_findings, total_tests_without_findings))
        is_last_test_case = True if test_case_number == total_reportable_tests else False
        if t.has_findings:
            if tests_with_findings == total_tests_with_findings or is_last_test_case:
                table = copy_table(document, table_with_findings, cut=True)
            else:
                table = copy_table(document,table_with_findings, cut=False)
        else:
            if tests_without_findings == total_tests_without_findings or is_last_test_case:
                table = copy_table(document, table_no_findings, cut=True)
            else:
                table = copy_table(document,table_no_findings, cut=False)

        # Classification Marking - Top
        cell = table.cell(0, 0)
        get_cleared_paragraph(cell).text = system_classification_verbose

        # Classification Marking - Bottom
        cell = table.cell(16, 1)
        get_cleared_paragraph(cell).text = system_classification_verbose

        # Test Case Number (Table Header Row)
        cell = table.cell(1, 0)
        if mission.test_case_identifier:
            get_cleared_paragraph(cell).text = 'Test #{0}-{1}'.format(mission.test_case_identifier, test_case_number)
        else:
            get_cleared_paragraph(cell).text = 'Test #{0}'.format(test_case_number)

        row_number = 0

        #
        # Test Case Title (Table Header Row)
        #
        row_number += 1
        cell = table.cell(row_number, 1)
        get_cleared_paragraph(cell).text = t.test_objective

        #
        # Attack Phase / Type
        #
        row_number += 1

        include_attack_phase = False
        include_attack_type = False

        if mission.attack_phase_include_flag \
                and t.attack_phase_include_flag \
                and len(t.get_attack_phase_display())> 0:
            include_attack_phase = True

        if mission.attack_type_include_flag \
                and t.attack_type_include_flag \
                and len(t.attack_type) > 0:
            include_attack_type = True

        left_cell = table.cell(row_number, 0)
        right_cell = get_cleared_paragraph(table.cell(row_number, 1))

        if include_attack_phase or include_attack_type:
            if include_attack_phase and include_attack_type:
                # Table text in column 1 assumes both items are included already
                right_cell.text = ' - '.join([t.get_attack_phase_display(), t.attack_type])
            elif include_attack_phase:
                get_cleared_paragraph(left_cell).text = "Attack Phase:"
                right_cell.text = t.get_attack_phase_display()
            elif include_attack_type:
                get_cleared_paragraph(left_cell).text = "Attack Type:"
                right_cell.text = t.attack_type
        else:
            logger.debug('Removing Attack Phase/Type Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Assumptions
        #
        row_number += 1
        if mission.assumptions_include_flag and t.assumptions_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.assumptions)
        else:
            logger.debug('Removing Assumptions Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Description
        #
        row_number += 1
        if mission.test_description_include_flag and t.test_description_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            if t.re_eval_test_case_number:
                cell.text = standardize_report_output_field('This is a reevaluation; reference previous test case #{0}\n\n{1}'.format(t.re_eval_test_case_number, t.test_description))
            else:
                cell.text = standardize_report_output_field(t.test_description)
        else:
            logger.debug('Removing Description Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Findings
        #
        row_number += 1
        if mission.findings_include_flag and t.findings_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.findings)
        else:
            logger.debug('Removing Findings Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Mitigations
        #
        row_number += 1
        if mission.mitigation_include_flag and t.mitigation_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.mitigation)
        else:
            logger.debug('Removing Mitigations Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Tools
        #
        row_number += 1
        if mission.tools_used_include_flag and t.tools_used_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.tools_used)
        else:
            logger.debug('Removing Tools Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Commands / Syntax
        #
        row_number += 1
        if mission.command_syntax_include_flag and t.command_syntax_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.command_syntax)
        else:
            logger.debug('Removing Commands/Syntax Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Targets
        #
        row_number += 1
        if mission.targets_include_flag and t.targets_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = '\n'.join([str(x) for x in t.target_hosts.all()])
        else:
            logger.debug('Removing Targets Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Sources
        #
        row_number += 1
        if mission.sources_include_flag and t.sources_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = '\n'.join([str(x) for x in t.source_hosts.all()])

        else:
            logger.debug('Removing Sources Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Date / Time
        #
        row_number += 1
        if mission.attack_time_date_include_flag and t.attack_time_date_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = localtime(t.attack_time_date).strftime('%b %d, %Y @ %I:%M %p')
        else:
            logger.debug('Removing Date/Time Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Side Effects
        #
        row_number += 1
        if mission.attack_side_effects_include_flag and t.attack_side_effects_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.attack_side_effects)
        else:
            logger.debug('Removing Side Effects Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Details
        #
        row_number += 1
        if mission.test_result_observation_include_flag and t.test_result_observation_include_flag:
            cell = get_cleared_paragraph(table.cell(row_number, 1))
            cell.text = standardize_report_output_field(t.test_result_observation)
        else:
            logger.debug('Removing Details Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Supporting Data
        #
        row_number += 1
        supporting_data_cell = None
        supporting_data_row = None
        if mission.supporting_data_include_flag:
            supporting_data_cell = get_cleared_paragraph(table.cell(row_number, 1))
            supporting_data_row = table.rows[row_number]
        else:
            logger.debug('Removing Supporting Data Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        #
        # Notes - Used for post report generation notes / customer use
        #
        row_number += 1
        if False: #TODO: add mission-level toggle
            logger.debug('Removing Customer Notes Row')
            remove_row(table, table.rows[row_number])
            row_number -= 1

        # Attachments section

        supporting_data_cell_items = []

        if mission.supporting_data_include_flag:
            my_data = TestSortingHelper.get_ordered_supporting_data(
                test_detail_id=t.id,
                reportable_supporting_data_only=True)

            if len(my_data) > 0:

                is_first_screenshot = True

                for d in my_data:
                    allowed_image_types = [
                        'gif',
                        'tiff',
                        'jpeg',
                        'bmp',
                        'png',
                    ]
                    try:
                        file_path = os.path.join(settings.MEDIA_ROOT, d.filename())
                        logger.debug('>> Beginning processing of {} at {}.'
                                     .format(
                                        d.filename(),
                                        file_path,
                                        ))

                        if imghdr.what(file_path) not in allowed_image_types:
                            raise UnrecognizedImageError('File type is not in the allowed image types. '
                                                         'Handling as non-image.')

                        if is_first_screenshot:
                            document.add_heading('Screenshots / Diagrams', level=3)
                            logger.debug('This is the first screenshot of this test case.')
                            is_first_screenshot = False

                        image_table = copy_table(document, data_table)
                        logger.debug('Creating a new image table.')
                        document.add_paragraph()

                        # Classification Marking - Top
                        cell = image_table.cell(0, 0)
                        get_cleared_paragraph(cell).text = system_classification_verbose

                        # Classification Marking - Bottom
                        cell = image_table.cell(2, 0)
                        get_cleared_paragraph(cell).text = system_classification_verbose

                        content_cell = image_table.cell(1, 0)

                        get_cleared_paragraph(content_cell).add_run().add_picture(d.test_file, width=Inches(5))
                        content_cell.paragraphs[0].add_run("\r" + d.caption)

                    except UnrecognizedImageError as e:
                        logger.debug('>> Attachment {attachment_name} not recognized as an image; adding as file.'
                                     .format(attachment_name=d.filename()))
                        supporting_data_cell_items.append('- {filename}: {caption}'.format(
                            filename=d.filename(),
                            caption=d.caption,
                        ))

                        if zip_attachments:
                            add_to_data_dir(mission_data_dir, test_case_number, d)
                    except (InvalidImageStreamError,
                            UnexpectedEndOfFileError) as e:
                        logger.warning('>> Attempting to add {file_name} to the report output resulted in an error: '
                                       '\n{trace}'.format(file_name=d.filename, trace=traceback.format_exc(10)))
                    except OSError as e:
                        logger.warning('>> Attempting to add {file_name} to the report output resulted in an error: '
                                       '\n{trace}'.format(file_name=d.filename, trace=traceback.format_exc(10)))
                        try:
                            if not d.test_file.closed:
                                d.test_file.close()
                        except IOError as e:
                            logger.warning('>> Attempting to close {file_name} resulted in an error: '
                                           '\n{trace}'.format(file_name=d.filename, trace=traceback.format_exc(10)))
                            pass

                if len(supporting_data_cell_items) > 0:
                    logger.debug('There are {} data cell items for TC {}.'.format(
                        len(supporting_data_cell_items),
                        t.id)
                    )
                    supporting_data_cell.text = '\n'.join(supporting_data_cell_items)

        if len(supporting_data_cell_items) == 0:
            logger.debug('There are no supporting_data_cell_items; removing the supporting data row.')
            remove_row(table, supporting_data_row)

    # Conclusion H1 and text
    document.add_heading('Conclusion', level=1)
    portion_mark_and_insert(mission.conclusion, document)

    data_table.cell(0, 0).text = ""
    get_cleared_paragraph(data_table.cell(1, 0)).text = "This table is used during report generation and can be deleted in the final report output."
    data_table.cell(2, 0).text = ""

    # Replace document slugs
    replace_document_slugs(document)

    if zip_attachments:
        zip_file = shutil.make_archive(mission_data_dir, 'zip', mission_data_dir)
        with open(mission_data_dir + '.zip', 'rb') as f:
            return io.BytesIO(f.read())

    else:
        my_stream = StringIO()
        document.save(my_stream)
        return my_stream
